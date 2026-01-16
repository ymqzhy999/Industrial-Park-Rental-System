import io
import zipfile
from flask import Blueprint, request, jsonify, send_file
from sqlalchemy import or_
from docx import Document
# 修正导入路径
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.models import FactoryInfo, User, Contract
from ..extensions import db

admin_bp = Blueprint('admin', __name__, url_prefix='/api/admin')


@admin_bp.route('/factory/update', methods=['POST'])
def update_factory():
    """更新厂房信息"""
    data = request.get_json()
    f_id = data.get('id')
    f = FactoryInfo.query.get(f_id)
    if not f: return jsonify({'code': 404, 'msg': '不存在'}), 404

    f.title = data.get('title')
    f.area = data.get('area')
    f.price_per_month = data.get('price')
    f.address = data.get('address')
    f.images = data.get('image')

    # 工业参数
    f.floor_height = data.get('floor_height')
    f.power_supply = data.get('power_supply')
    f.floor_load = data.get('floor_load')
    f.structure = data.get('structure')

    # 处理标签
    tags_data = data.get('tags')
    if isinstance(tags_data, list):
        f.tags = ','.join(tags_data)
    else:
        f.tags = tags_data

    db.session.commit()
    return jsonify({'code': 200, 'msg': '更新成功'})


@admin_bp.route('/contract/list', methods=['POST'])
def list_contracts():
    """管理员获取合同列表"""
    data = request.get_json() or {}
    page = data.get('page', 1)
    page_size = data.get('pageSize', 10)
    keyword = data.get('keyword', '')
    status = data.get('status')

    query = Contract.query.outerjoin(User)

    if keyword:
        rule = or_(
            Contract.contract_no.ilike(f"%{keyword}%"),
            Contract.party_b.ilike(f"%{keyword}%"),
            User.real_name.ilike(f"%{keyword}%")
        )
        query = query.filter(rule)

    if status is not None and status != '':
        query = query.filter(Contract.status == int(status))

    pagination = query.order_by(Contract.create_time.desc()).paginate(page=page, per_page=page_size, error_out=False)

    res = []
    for c in pagination.items:
        res.append({
            'id': c.contract_id,
            'no': c.contract_no,
            'tenant_name': c.party_b or (c.user.real_name if c.user else '未知'),
            'factory_title': c.factory.title if c.factory else '已删除',
            'total_price': c.total_price,
            'status': c.status,
            'create_time': c.create_time.strftime('%Y-%m-%d')
        })
    return jsonify({'code': 200, 'data': res, 'total': pagination.total})


@admin_bp.route('/contract/batch_export', methods=['POST'])
def batch_export_contracts():
    """批量导出合同 (管理员)"""
    data = request.get_json()
    ids = data.get('ids', [])

    if not ids:
        return jsonify({'code': 400, 'msg': '未选择任何合同'}), 400

    contracts = Contract.query.filter(Contract.contract_id.in_(ids)).all()
    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        for c in contracts:
            tenant_name = c.party_b or (c.user.real_name if c.user else '未知用户')

            doc = Document()
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

            head = doc.add_heading(f'合同-{tenant_name}', 0)
            head.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f'合同编号：{c.contract_no}').alignment = WD_ALIGN_PARAGRAPH.RIGHT

            p = doc.add_paragraph()
            p.add_run(f"甲方：{c.party_a or '园区管理方'}\n")
            p.add_run(f"乙方：{tenant_name}\n")
            p.add_run(f"租赁标的：{c.factory.title if c.factory else '未知厂房'}\n")
            p.add_run(f"金额：{c.total_price} 元\n")
            p.add_run(f"租期：{c.start_date} 至 {c.end_date}")

            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            filename = f"{c.contract_no}_{tenant_name}.docx"
            zf.writestr(filename, docx_buffer.getvalue())

    zip_buffer.seek(0)
    return send_file(
        zip_buffer,
        mimetype='application/zip',
        as_attachment=True,
        download_name='contracts_export.zip'
    )


@admin_bp.route('/stats', methods=['GET'])
def get_stats():
    """获取管理后台首页统计数据"""
    user_count = User.query.filter_by(role=1).count()
    factory_count = FactoryInfo.query.count()
    pending_contract = Contract.query.filter_by(status=0).count()
    total_income = db.session.query(db.func.sum(Contract.total_price)).filter_by(status=1).scalar() or 0

    return jsonify({
        'code': 200,
        'data': {
            'user_count': user_count,
            'factory_count': factory_count,
            'pending_contract': pending_contract,
            'total_income': total_income
        }
    })


@admin_bp.route('/factory/list', methods=['POST'])
def list_factory():
    """管理员获取厂房列表"""
    data = request.get_json() or {}
    page = data.get('page', 1)
    page_size = data.get('pageSize', 10)
    keyword = data.get('keyword', '')

    query = FactoryInfo.query
    if keyword:
        query = query.filter(FactoryInfo.title.ilike(f"%{keyword}%"))

    pagination = query.order_by(FactoryInfo.factory_id.desc()).paginate(page=page, per_page=page_size, error_out=False)

    res = []
    for f in pagination.items:
        res.append({
            'id': f.factory_id,
            'title': f.title,
            'area': f.area,
            'price': f.price_per_month,
            'status': f.status,
            'image': f.images.split(',')[0] if f.images else '',
            'address': f.address
        })

    return jsonify({'code': 200, 'data': res, 'total': pagination.total})


@admin_bp.route('/factory/add', methods=['POST'])
def add_factory():
    """添加厂房"""
    data = request.get_json()
    new_f = FactoryInfo(
        title=data.get('title'),
        area=data.get('area'),
        price_per_month=data.get('price'),
        address=data.get('address'),
        floor_height=data.get('floor_height'),
        power_supply=data.get('power_supply'),
        images=data.get('image', ''),
        status=0
    )
    db.session.add(new_f)
    db.session.commit()
    return jsonify({'code': 200, 'msg': '发布成功'})


@admin_bp.route('/factory/delete', methods=['POST'])
def delete_factory():
    """删除厂房"""
    data = request.get_json()
    f_id = data.get('id')
    f = FactoryInfo.query.get(f_id)
    if not f: return jsonify({'code': 404, 'msg': '不存在'}), 404

    try:
        db.session.delete(f)
        db.session.commit()
        return jsonify({'code': 200, 'msg': '删除成功'})
    except Exception:
        db.session.rollback()
        return jsonify({'code': 500, 'msg': '该厂房存在关联合同，无法直接删除，请先处理合同'}), 500


@admin_bp.route('/user/list', methods=['GET'])
def list_users():
    """获取用户列表"""
    users = User.query.all()
    res = []
    for u in users:
        res.append({
            'id': u.user_id,
            'username': u.username,
            'real_name': u.real_name,
            'role': u.role,
            'create_time': u.create_time.strftime('%Y-%m-%d')
        })
    return jsonify({'code': 200, 'data': res})