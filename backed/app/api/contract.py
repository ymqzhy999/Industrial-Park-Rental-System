import io
import datetime
import random
import zipfile
from flask import Blueprint, request, jsonify, send_file
from sqlalchemy import desc
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn

from app.models import Contract, FactoryInfo, User
from app.extensions import db

contract_bp = Blueprint('contract', __name__, url_prefix='/api/contract')


def generate_contract_docx(contract):
    """
    生成合同 Word 文档流
    :param contract: Contract 对象
    :return: BytesIO 对象 (Word 文件流)
    """
    doc = Document()

    # 1. 设置中文字体兼容性
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    style.font.size = Pt(12)

    # 2. 标题
    heading = doc.add_heading('工业园区厂房租赁合同', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3. 基本信息
    sign_date = contract.create_time.strftime("%Y-%m-%d") if contract.create_time else datetime.datetime.now().strftime(
        "%Y-%m-%d")

    doc.add_paragraph(f'合同编号：{contract.contract_no}')
    doc.add_paragraph(f'甲方（出租方）：{contract.party_a or "智慧工业园区管理中心"}')
    doc.add_paragraph(f'乙方（承租方）：{contract.party_b}')
    doc.add_paragraph(f'签订日期：{sign_date}')
    doc.add_paragraph('-' * 30)

    # 4. 正文内容
    content = contract.content
    if not content:
        content = f"（暂无详细条款内容）\n租赁厂房：{contract.factory.title if contract.factory else '未知'}\n总金额：{contract.total_price}元"

    for line in content.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line.strip())
            p.paragraph_format.space_after = Pt(6)

    # 5. 保存到内存流
    byte_io = io.BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io


@contract_bp.route('/apply', methods=['POST'])
def apply_contract():
    """提交租赁申请"""
    data = request.get_json()
    username = data.get('username')
    factory_id = data.get('factory_id')
    start_date_str = data.get('start_date')
    end_date_str = data.get('end_date')
    purpose = data.get('purpose')
    contact_phone = data.get('phone')
    terms = data.get('terms', '无')

    if not all([username, factory_id, start_date_str, end_date_str]):
        return jsonify({'code': 400, 'msg': '参数不完整'}), 400

    user = User.query.filter_by(username=username).first()
    factory = FactoryInfo.query.get(factory_id)

    if not user or not factory:
        return jsonify({'code': 404, 'msg': '用户或厂房不存在'}), 404

    if factory.status != 0:
        return jsonify({'code': 400, 'msg': '该厂房已被租赁，无法申请'}), 400

    # 计算总价
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%Y-%m-%d')
        end_date = datetime.datetime.strptime(end_date_str, '%Y-%m-%d')
        days = (end_date - start_date).days
        total_price = round((days / 30) * factory.price_per_month, 2)
    except Exception as e:
        print(f"Error calculating price: {e}")
        start_date = datetime.datetime.now()
        end_date = start_date
        total_price = factory.price_per_month

    # 生成编号
    date_str = datetime.datetime.now().strftime('%Y%m%d')
    rand_str = str(random.randint(1000, 9999))
    contract_no = f"CTR{date_str}{rand_str}"

    # 生成合同正文
    contract_content = (
        f"【工业园区厂房租赁合同】\n\n"
        f"合同编号：{contract_no}\n"
        f"出租方（甲方）：智慧工业园区管理中心\n"
        f"承租方（乙方）：{user.real_name or username}\n"
        f"联系电话：{contact_phone}\n\n"
        f"一、租赁标的\n"
        f"甲方同意将位于 {factory.address} 的厂房（{factory.title}）出租给乙方使用，建筑面积 {factory.area} 平方米。\n\n"
        f"二、租赁期限\n"
        f"自 {start_date_str} 起至 {end_date_str} 止。\n\n"
        f"三、租赁费用\n"
        f"每月租金人民币 {factory.price_per_month} 元。\n"
        f"合同期内预计总租金：{total_price} 元。\n\n"
        f"四、租赁用途\n"
        f"乙方承诺租赁该厂房用于：{purpose}。\n\n"
        f"五、补充条款（由乙方申请时提出）\n"
        f"{terms}\n\n"
        f"六、违约责任\n"
        f"双方应严格遵守本合同，任何一方违约需承担法律责任。\n\n"
        f"签署日期：{datetime.datetime.now().strftime('%Y年%m月%d日')}"
    )

    new_contract = Contract(
        contract_no=contract_no,
        user_id=user.user_id,
        factory_id=factory.factory_id,
        start_date=start_date,
        end_date=end_date,
        total_price=total_price,
        status=0,
        party_a="园区管理方",
        party_b=user.real_name or username,
        content=contract_content
    )

    try:
        db.session.add(new_contract)
        db.session.commit()
        return jsonify({'code': 200, 'msg': '申请已提交，等待审核'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'code': 500, 'msg': '保存失败'}), 500


@contract_bp.route('/list', methods=['POST'])
def get_contract_list():
    """获取合同列表 (支持分页 + 搜索 + Tab过滤)"""
    data = request.get_json() or {}
    page = data.get('page', 1)
    page_size = data.get('pageSize', 10)
    username = data.get('username')
    tab = data.get('tab', 'active')

    query = Contract.query

    # 1. 过滤用户
    if username:
        user = User.query.filter_by(username=username).first()
        if user and user.role != 0:
            query = query.filter_by(user_id=user.user_id)

    # 2. 状态过滤
    if tab == 'active':
        query = query.filter(Contract.status == 1)
    elif tab == 'pending':
        query = query.filter(Contract.status.in_([0, 3]))
    elif tab == 'archived':
        query = query.filter(Contract.status.in_([2, 4]))

    # 3. 分页
    pagination = query.order_by(desc(Contract.create_time)).paginate(page=page, per_page=page_size, error_out=False)

    def safe_date(d):
        if not d: return ''
        if isinstance(d, str): return d.split(' ')[0]
        try:
            return d.strftime('%Y-%m-%d')
        except:
            return str(d).split(' ')[0]

    res_data = []
    for c in pagination.items:
        res_data.append({
            'id': c.contract_id,
            'no': c.contract_no,
            'factory_title': c.factory.title if c.factory else '厂房失效',
            'tenant_name': c.user.real_name if c.user else '未知用户',
            'start_date': safe_date(c.start_date),
            'end_date': safe_date(c.end_date),
            'price': c.total_price,
            'status': c.status,
            'total_price': c.total_price
        })

    return jsonify({
        'code': 200,
        'data': res_data,
        'total': pagination.total
    })


@contract_bp.route('/detail/<int:contract_id>', methods=['GET'])
def get_contract_detail(contract_id):
    """获取合同详情"""
    c = Contract.query.get(contract_id)
    if not c:
        return jsonify({'code': 404, 'msg': '合同不存在'}), 404

    def safe_date(d):
        if not d: return ''
        if isinstance(d, str): return d.split(' ')[0]
        try:
            return d.strftime('%Y-%m-%d')
        except:
            return str(d).split(' ')[0]

    return jsonify({
        'code': 200,
        'data': {
            'id': c.contract_id,
            'contract_no': c.contract_no,
            'party_a': c.party_a,
            'party_b': c.party_b,
            'factory_title': c.factory.title if c.factory else '未知厂房',
            'factory_address': c.factory.address if c.factory else '',
            'start_date': safe_date(c.start_date),
            'end_date': safe_date(c.end_date),
            'create_time': safe_date(c.create_time),
            'total_price': c.total_price,
            'status': c.status,
            'content': c.content
        }
    })


@contract_bp.route('/export/<int:contract_id>', methods=['GET'])
def export_contract(contract_id):
    """导出单个合同 (Word)"""
    c = Contract.query.get(contract_id)
    if not c: return "合同不存在", 404

    file_stream = generate_contract_docx(c)
    filename = f"合同_{c.party_b}_{c.contract_no}.docx"

    try:
        filename = filename.encode('latin-1').decode('latin-1')
    except:
        pass

    return send_file(
        file_stream,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@contract_bp.route('/export/batch', methods=['POST'])
def export_batch_contracts():
    """批量导出合同 (ZIP)"""
    data = request.get_json()
    ids = data.get('ids', [])

    if not ids:
        return jsonify({'code': 400, 'msg': '未选择合同'}), 400

    contracts = Contract.query.filter(Contract.contract_id.in_(ids)).all()
    if not contracts:
        return jsonify({'code': 404, 'msg': '未找到有效合同'}), 404

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for c in contracts:
            docx_stream = generate_contract_docx(c)
            filename = f"{c.party_b}_{c.contract_no}.docx"
            zip_file.writestr(filename, docx_stream.getvalue())

    zip_buffer.seek(0)
    return send_file(
        zip_buffer,
        as_attachment=True,
        download_name=f"批量导出合同_{datetime.datetime.now().strftime('%Y%m%d')}.zip",
        mimetype='application/zip'
    )


@contract_bp.route('/delete', methods=['POST'])
def delete_contracts():
    """批量删除合同"""
    data = request.get_json()
    ids = data.get('ids', [])

    if not ids:
        return jsonify({'code': 400, 'msg': '请选择要删除的合同'}), 400

    try:
        Contract.query.filter(Contract.contract_id.in_(ids)).delete(synchronize_session=False)
        db.session.commit()
        return jsonify({'code': 200, 'msg': '删除成功'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'code': 500, 'msg': str(e)}), 500


@contract_bp.route('/audit', methods=['POST'])
def audit_contract():
    """审核合同"""
    data = request.get_json()
    c_id = data.get('id')
    action = data.get('action')

    contract = Contract.query.get(c_id)
    if not contract: return jsonify({'code': 404, 'msg': '合同不存在'}), 404

    if action == 'pass':
        contract.status = 3  # 待支付
    elif action == 'reject':
        contract.status = 2  # 已拒绝

    db.session.commit()
    return jsonify({'code': 200, 'msg': '操作成功'})